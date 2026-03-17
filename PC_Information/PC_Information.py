# pip install python-docx

import time
from datetime import datetime as dt
from platform import node
from wmi import WMI

from docx import Document


def get_size(bts: int, ending='iB') -> str:
    size = 1024
    for item in ["", "K", "M", "G", "T", "P"]:
        if bts < size:
            return f"{bts:.2f} {item}{ending}" if bts > 0 else f"{bts:.2f} {item}B"
        bts /= size


# OS
def wmi_os() -> dict:
    os_info = dict()

    c = WMI()
    u_name = "Unknown"
    if cs := c.Win32_ComputerSystem()[0]:
        os_info.update({"ComputerName": cs.Name})
        u_name = [user.UserName for user in c.Win32_ComputerSystem()]
    if cs := c.Win32_OperatingSystem()[0]:
        os_info.update({
            'Caption': cs.Caption,
            "LastBootUpTime": dt.strptime(cs.LastBootUpTime[:14], '%Y%m%d%H%M%S').strftime("%Y-%m-%d %H:%M:%S"),
            'Version': cs.Version,
            'BuildNumber': cs.BuildNumber,
            'InstallDate': dt.strptime(cs.InstallDate[:14], '%Y%m%d%H%M%S').strftime("%Y-%m-%d %H:%M:%S"),
            'WindowsDirectory': cs.WindowsDirectory,
            'UserName': u_name
        })
    return os_info if os_info else False


# BIOS
def bios_wmi() -> (dict, bool):
    bios_info = dict()
    if c := WMI().Win32_BIOS()[0]:
        bios_info.update({
            'Manufacturer': c.Manufacturer,
            'Name': c.Name,
            'Version': c.Version
        })
    return bios_info if bios_info else False


# Motherboard
def motherboard_wmi() -> (dict, bool):
    motherboard = dict()
    if c := WMI().Win32_BaseBoard()[0]:
        motherboard.update({
            'Manufacturer': c.Manufacturer,
            'Product': c.Product
        })
    return motherboard if motherboard else False


# CPU
def cpu_wmi() -> (dict, bool):
    cpu_info = dict()
    w = WMI()
    if cpu := w.Win32_Processor()[0]:
        cpu_info.update({
            'Name': cpu.Name,
            'ID': cpu.ProcessorId,
            'NumberOfLogicalProcessors': cpu.NumberOfLogicalProcessors,
            'NumberOfPhysicalProcessors': cpu.NumberOfCores,
            'Manufacturer': cpu.Manufacturer,
            'MaxClockSpeed': f"{cpu.MaxClockSpeed} MHz",
            'Socket': cpu.SocketDesignation,
            'Caption': cpu.Caption
        })
    return cpu_info if cpu_info else False


# GPU
def gpu_wmi() -> (dict, bool):
    gpu = dict()
    if vc := WMI().Win32_VideoController()[0]:
        gpu.update({
            "Name": vc.Description,
            "AdapterRAM": get_size(abs(vc.AdapterRAM)),
            "Resolution": f'{vc.CurrentHorizontalResolution}x{vc.CurrentVerticalResolution}',
            "CurrentRefreshRate": f'{vc.CurrentRefreshRate} Гц',
            "VideoProcessor": vc.VideoProcessor
        })
    return gpu if gpu else False


# Memory
def memory_wmi() -> (dict, bool):
    memory_dict = dict()
    wmi_obj = WMI()
    if memory_data := wmi_obj.Win32_PhysicalMemory():
        memory_dict.update({"TotalPhysicalMemory": get_size(sum([int(mem.Capacity) for mem in memory_data]))})
        for i, mem in enumerate(memory_data):
            memory_dict[f"Physical Memory {i}"] = dict()
            memory_dict[f"Physical Memory {i}"].update({
                "Capacity": get_size(int(mem.Capacity)),
                "ConfiguredClockSpeed": mem.Speed,
                "Manufacturer": mem.Manufacturer,
                "PartNumber": mem.PartNumber,
                "SerialNumber": mem.SerialNumber
            })
    return memory_dict if memory_dict else False


# HDD, SSD
def hdd_ssd_wmi() -> (dict, bool):
    disk_info = dict()
    c = WMI()
    if disks := c.Win32_DiskDrive():
        for disk in disks:
            disk_info[disk.DeviceID] = {
                'Caption': disk.Model,
                'MediaType': disk.InterfaceType,
                'Capacity': get_size(int(disk.Size)) if disk.Size else 'None'
            }
    return disk_info if disk_info else False


# CD-ROM
def cdrom_wmi() -> (dict, bool):
    cdrom_info = dict()
    c = WMI()
    if cdroms := c.Win32_CDROMDrive():
        for cdrom in cdroms:
            cdrom_info[cdrom.Caption] = dict()
            cdrom_info[cdrom.Caption].update({
                'Drive': cdrom.Caption,
                'MediaType': cdrom.MediaType,
                'Status': cdrom.Status,
                'SerialNumber': cdrom.SerialNumber,
                'Manufacturer': cdrom.Manufacturer
            })
    return cdrom_info if cdrom_info else False


# Network Interface
def nic_wmi() -> (dict, bool):
    nic = dict()
    conn = WMI()
    description = [it.Description for it in conn.Win32_NetworkAdapter() if it.PhysicalAdapter]
    if description:
        for it in conn.Win32_NetworkAdapterConfiguration():
            if it.Description in description:
                desc = it.Description
                nic[desc] = dict()
                try:
                    nic[desc].update({"DefaultIPGateway": it.DefaultIPGateway[0]})
                except TypeError:
                    nic[desc].update({"DefaultIPGateway": None})
                try:
                    nic[desc].update({"DHCPServer": it.DHCPServer})
                except TypeError:
                    nic[desc].update({"DHCPServer": None})
                try:
                    nic[desc].update({"DNSHostName": it.DNSHostName})
                except TypeError:
                    nic[desc].update({"DNSHostName": None})
                try:
                    nic[desc].update({"IPv4Address": it.IPAddress[0]})
                except TypeError:
                    nic[desc].update({"IPv4Address": None})
                try:
                    nic[desc].update({"IPv6Address": it.IPAddress[1]})
                except TypeError:
                    nic[desc].update({"IPv6Address": None})
                try:
                    nic[desc].update({"IPSubnet": it.IPSubnet[0]})
                except TypeError:
                    nic[desc].update({"IPSubnet": None})
                try:
                    nic[desc].update({"MACAddress": it.MACAddress})
                except TypeError:
                    nic[desc].update({"MACAddress": None})
                try:
                    nic[desc].update({"ServiceName": it.ServiceName})
                except TypeError:
                    nic[desc].update({"ServiceName": None})
    return nic if nic else False


wmic_info = ""


def print_wmic(part, dict_info):
    global wmic_info
    synonyms = {"ComputerName": "Имя компьютера", "Caption": "Название", "InstallDate": "Дата установки",
                "LastBootUpTime": "Время последней загрузки", "Version": "Версия",
                "WindowsDirectory": "Директория Windows", "TimeZone": "Часовой пояс", "UserName": "Имя пользователя",
                "Manufacturer": "Производитель", "Name": "Название", "Product": "Изделие",
                "MaxClockSpeed": "Максимальная тактовая частота", "SocketDesignation": "Название сокета",
                "NumberOfPhysicalProcessors": "Количество физических процессоров", "VideoProcessor": "Видеопроцессор",
                "NumberOfLogicalProcessors": "Количество логических процессоров", "Capacity": "Емкость",
                "AdapterRAM": "Оперативная память адаптера", "CurrentRefreshRate": "Текущая частота обновления",
                "Resolution": "Разрешение", "TotalPhysicalMemory": "Общий объем физической памяти", "Socket": "Сокет",
                "ConfiguredClockSpeed": "Настроенная тактовая частота", "PartNumber": "Номер партии",
                "SerialNumber": "Серийный номер", "DeviceID": "Идентификатор устройства", "MediaType": "Тип носителя",
                "FirmwareRevision": "Ревизия прошивки", "Partitions": "Разделы", "Size": "Объем", "Drive": "Диск",
                "VolumeName": "Имя тома", "VolumeSerialNumber": "Серийный номер тома", "MACAddress": "MAC-адрес",
                "NetConnectionID": "Идентификатор сетевого подключения", "DHCPServer": "DHCP-сервер",
                "IPAddress": "IP-адрес", "BuildNumber": "Номер сборки", "ID": "Идентификатор", "Status": "Статус",
                "DefaultIPGateway": "IP-адрес шлюза по-умолчанию", "DNSHostName": "DNS Имя хоста",
                "IPv4Address": "IPv4-адрес", "IPv6Address": "IPv6-адрес", "IPSubnet": "Маска подсети",
                "ServiceName": "Название службы"}
    part += f'{"-" * 50}\n'
    for key in dict_info:
        if type(dict_info[key]) == dict:
            for item in dict_info[key]:
                part += f'{synonyms[item]}: {dict_info[key][item]}\n'
            part += "\n"
        else:
            part += f'{synonyms[key]}: {dict_info[key]}\n'
    print(part)
    wmic_info += f'{part}\n'


def main():
    global wmic_info
    t = time.monotonic()
    document = Document()
    document.add_heading(f'Сводная информация о компьютере: {node()}', 0)

    if os_info := wmi_os():
        print_wmic("Информация об операционной системе\n", os_info)
    if bios_info := bios_wmi():
        print_wmic("Информация о BIOS\n", bios_info)
    if mb_info := motherboard_wmi():
        print_wmic("Информация о материнской плате\n", mb_info)
    if cpu_info := cpu_wmi():
        print_wmic("Информация о процессоре\n", cpu_info)
    if gpu_info := gpu_wmi():
        print_wmic("Информация о видеокарте\n", gpu_info)
    if mem_info := memory_wmi():
        print_wmic("Информация об оперативной памяти\n", mem_info)
    if drive_info := hdd_ssd_wmi():
        print_wmic("Информация о HDD и SSD\n", drive_info)
    if cd_rom_info := cdrom_wmi():
        print_wmic("Информация о CD/DVD-ROM\n", cd_rom_info)
    if nic_info := nic_wmi():
        print_wmic("Информация о физических сетевых интерфейсах\n", nic_info)

    document.add_paragraph(wmic_info)
    document.save(f'{node()}.docx')
    print(f"Собранная информация сохранена в файл: {node()}.docx")
    print(f'\nВремя работы скрипта: {time.monotonic() - t} с.')


if __name__ == "__main__":
    main()